from __future__ import annotations

import sys
from pathlib import Path

import pytesseract
from docx import Document
from PIL import Image, ImageDraw

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from src.extractor import extract_text


SAMPLES_DIR = Path(__file__).parent / "samples"


def _create_sample_docx(path: Path) -> None:
    document = Document()
    document.add_heading("Sample DOCX", level=1)
    document.add_paragraph("Bonjour depuis DOCX")
    document.add_paragraph("Montant: 123.45 EUR")
    document.save(path)


def _create_sample_image(path: Path) -> None:
    image = Image.new("RGB", (900, 240), color="white")
    draw = ImageDraw.Draw(image)
    draw.text((20, 30), "OCR SAMPLE", fill="black")
    draw.text((20, 120), "Invoice 2026-07-12", fill="black")
    image.save(path)


def _create_sample_pdf(path: Path) -> None:
    # Minimal PDF with embedded text stream for smoke extraction tests.
    pdf_bytes = b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>
endobj
4 0 obj
<< /Length 58 >>
stream
BT
/F1 24 Tf
72 720 Td
(Bonjour PDF exemple) Tj
ET
endstream
endobj
5 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj
xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000241 00000 n 
0000000350 00000 n 
trailer
<< /Root 1 0 R /Size 6 >>
startxref
420
%%EOF
"""
    path.write_bytes(pdf_bytes)


def _check_contains(label: str, text: str, expected: str) -> bool:
    ok = expected.lower() in text.lower()
    status = "OK" if ok else "KO"
    print(f"[{status}] {label}: contient '{expected}'")
    return ok


def main() -> int:
    SAMPLES_DIR.mkdir(parents=True, exist_ok=True)

    sample_docx = SAMPLES_DIR / "sample.docx"
    sample_pdf = SAMPLES_DIR / "sample.pdf"
    sample_image = SAMPLES_DIR / "sample.png"

    _create_sample_docx(sample_docx)
    _create_sample_pdf(sample_pdf)
    _create_sample_image(sample_image)

    print("=== Round 3: extraction PDF / DOCX / OCR ===")

    all_ok = True

    docx_text = extract_text(sample_docx)
    all_ok = _check_contains("DOCX", docx_text, "Bonjour depuis DOCX") and all_ok

    pdf_text = extract_text(sample_pdf)
    all_ok = _check_contains("PDF", pdf_text, "Bonjour PDF exemple") and all_ok

    try:
        ocr_text = extract_text(sample_image)
        all_ok = _check_contains("OCR", ocr_text, "OCR") and all_ok
    except pytesseract.TesseractNotFoundError:
        print("[SKIP] OCR: Tesseract non installe ou non configure (TESSERACT_CMD).")

    if all_ok:
        print("Round 3 OK: extraction PDF/DOCX/OCR validee.")
        return 0

    print("Round 3 KO: au moins une extraction n'a pas retourne le texte attendu.")
    return 1


if __name__ == "__main__":
    raise SystemExit(main())
